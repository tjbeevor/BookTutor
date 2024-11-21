import streamlit as st
import google.generativeai as genai
import PyPDF2
import docx
import markdown
import bs4
import io
import json
import time
import random
import re
from typing import Dict, List, Optional, Any, BinaryIO, Union
from dataclasses import dataclass
from abc import ABC, abstractmethod
from pathlib import Path
from dotenv import load_dotenv
import os

# Load environment variables
load_dotenv()

# Utility functions

def extract_json_from_text(text: str) -> str:
    """Extract JSON content from text using a more robust approach"""
    # Find the first { and last } in the text
    start = text.find('{')
    end = text.rfind('}')
    
    if start == -1 or end == -1:
        raise ValueError("No JSON structure found in text")
        
    # Extract the potential JSON content
    json_str = text[start:end + 1]
    return json_str


def clean_json_string(text: str) -> str:
    """Clean and extract JSON string from text"""
    json_match = re.search(r'\{[\s\S]*\}', text)
    if not json_match:
        raise ValueError("No JSON structure found in text")
        
    json_str = json_match.group()
    
    # Remove any markdown backticks or json indicators
    json_str = re.sub(r'^```json\s*', '', json_str)
    json_str = re.sub(r'\s*```$', '', json_str)
    
    # Fix common JSON formatting issues without using look-behind
    # Replace quotes after numbers using capture groups
    json_str = re.sub(r'(\d)"([\s,}])', r'\1\2', json_str)
    
    # Replace quotes after booleans using capture groups
    json_str = re.sub(r'(true|false)"([\s,}])', r'\1\2', json_str)
    
    # Remove trailing commas
    json_str = re.sub(r',\s*}', '}', json_str)
    
    return json_str

def create_fallback_evaluation_json() -> str:
    """Create a fallback JSON string for when parsing fails"""
    fallback = {
        "points_covered": [],
        "missing_points": [],
        "misconceptions": [],
        "understanding_level": 50,
        "strengths": ["Attempted to answer the question"],
        "areas_for_improvement": ["Clarity and completeness of response"],
        "suggestions": ["Please try to be more specific in your answer"]
    }
    return json.dumps(fallback)

def is_header_line(line: str) -> bool:
    """Detect if a line is likely a header"""
    line = line.strip()
    if not line:
        return False
        
    # Check for common header patterns
    header_indicators = [
        line.isupper(),
        line.startswith('#'),
        len(line.split()) <= 5 and line[0].isupper(),
        any(line.startswith(str(i) + '.') for i in range(1, 100))
    ]
    
    return any(header_indicators)

# Custom Exceptions
class DocumentProcessingError(Exception):
    """Custom exception for document processing errors"""
    pass

class UnsupportedFormatError(Exception):
    """Custom exception for unsupported file formats"""
    pass

class ValidationError(Exception):
    """Custom exception for content validation errors"""
    pass

# Base Classes and Data Classes
@dataclass
class Topic:
    title: str
    content: str
    subtopics: List['Topic']
    completed: bool = False
    parent: Optional['Topic'] = None
    metadata: Dict[str, Any] = None
    content_type: str = "general"
    difficulty_level: str = "intermediate"

    def __post_init__(self):
        if self.metadata is None:
            self.metadata = {}
            
    def add_subtopic(self, subtopic: 'Topic') -> None:
        """Add a subtopic and set its parent"""
        subtopic.parent = self
        self.subtopics.append(subtopic)
    
    def get_all_subtopics(self) -> List['Topic']:
        """Get all subtopics recursively"""
        all_subtopics = []
        for subtopic in self.subtopics:
            all_subtopics.append(subtopic)
            all_subtopics.extend(subtopic.get_all_subtopics())
        return all_subtopics

@dataclass
class UserProgress:
    understanding_level: float = 50.0
    completed_topics: List[str] = None
    current_approach: str = "practical"
    difficulty_level: str = "beginner"
    last_evaluation_date: Optional[float] = None
    
    def __post_init__(self):
        if self.completed_topics is None:
            self.completed_topics = []
        if self.last_evaluation_date is None:
            self.last_evaluation_date = time.time()
    
    def update_progress(self, topic_title: str, understanding_level: float) -> None:
        """Update progress with new evaluation data"""
        self.understanding_level = understanding_level
        if understanding_level >= 70 and topic_title not in self.completed_topics:
            self.completed_topics.append(topic_title)
        self.last_evaluation_date = time.time()
    
    def adjust_difficulty(self) -> None:
        """Adjust difficulty based on understanding level"""
        if self.understanding_level >= 85:
            self.difficulty_level = "advanced"
        elif self.understanding_level >= 65:
            self.difficulty_level = "intermediate"
        else:
            self.difficulty_level = "beginner"

class BaseDocumentProcessor(ABC):
    def __init__(self):
        self.supported_extensions: List[str] = []
        
    @abstractmethod
    def extract_content(self, file_obj: BinaryIO) -> Dict[str, Any]:
        """Extract content from the document"""
        pass

    @abstractmethod
    def validate_content(self, content: Dict[str, Any]) -> bool:
        """Validate extracted content"""
        pass

    def standardize_content(self, content: Dict[str, Any]) -> Dict[str, Any]:
        """Convert content to standard internal format"""
        return {
            'text': content.get('text', ''),
            'metadata': content.get('metadata', {}),
            'structure': content.get('structure', {}),
            'type': content.get('type', 'general'),
            'processed_date': time.time()
        }
    
    def process(self, file_obj: BinaryIO) -> Dict[str, Any]:
        """Complete processing pipeline for a document"""
        try:
            # Extract content
            content = self.extract_content(file_obj)
            
            # Validate
            if not self.validate_content(content):
                raise ValidationError("Invalid document content")
                
            # Standardize
            return self.standardize_content(content)
            
        except Exception as e:
            raise DocumentProcessingError(f"Processing error: {str(e)}")

# Document Processor Implementation
class PDFProcessor(BaseDocumentProcessor):
    def __init__(self):
        super().__init__()
        self.supported_extensions = ['.pdf']
    
    def extract_content(self, file_obj: BinaryIO) -> Dict[str, Any]:
        """Extract content from PDF files"""
        try:
            content = {
                'text': [],
                'metadata': {},
                'structure': {'sections': []},
                'type': 'pdf'
            }
            
            pdf_reader = PyPDF2.PdfReader(file_obj)
            
            # Extract metadata
            content['metadata'] = {
                'pages': len(pdf_reader.pages),
                'title': pdf_reader.metadata.get('/Title', 'Untitled'),
                'author': pdf_reader.metadata.get('/Author', 'Unknown'),
                'subject': pdf_reader.metadata.get('/Subject', ''),
                'creation_date': pdf_reader.metadata.get('/CreationDate', ''),
                'extract_date': time.strftime('%Y-%m-%d %H:%M:%S')
            }
            
            # Process each page
            current_section = {'title': None, 'content': []}
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if not page_text:
                        continue
                        
                    # Split into lines for analysis
                    lines = page_text.split('\n')
                    for line in lines:
                        line = line.strip()
                        if not line:
                            continue
                            
                        # Use utility function to detect headers
                        if is_header_line(line):
                            if current_section['title'] and current_section['content']:
                                content['structure']['sections'].append(current_section.copy())
                            current_section = {'title': line, 'content': []}
                        else:
                            current_section['content'].append(line)
                            
                    content['text'].append(page_text)
                    
                except Exception as e:
                    st.warning(f"Warning: Error processing page {page_num}: {str(e)}")
                    continue
            
            # Add the last section if it exists
            if current_section['title'] and current_section['content']:
                content['structure']['sections'].append(current_section)
                
            return content
            
        except Exception as e:
            raise DocumentProcessingError(f"Error processing PDF: {str(e)}")

    def validate_content(self, content: Dict[str, Any]) -> bool:
        """Validate PDF content"""
        try:
            # Check for required fields
            if not content.get('text'):
                return False
                
            # Check for non-empty text content
            if not any(text.strip() for text in content['text']):
                return False
                
            # Check for basic structure
            if not content.get('structure', {}).get('sections'):
                return False
                
            # Check metadata
            if not content.get('metadata', {}).get('pages'):
                return False
                
            return True
            
        except Exception as e:
            st.warning(f"Validation error: {str(e)}")
            return False

#Hello world

class DocxProcessor(BaseDocumentProcessor):
    def extract_content(self, file_obj: BinaryIO) -> Dict[str, Any]:
        """Extract content from DOCX files"""
        try:
            doc = docx.Document(file_obj)
            content = {
                'text': [],
                'metadata': {},
                'structure': {'sections': []},
                'type': 'docx'
            }
            
            # Extract document properties
            content['metadata'] = {
                'title': doc.core_properties.title or 'Untitled',
                'author': doc.core_properties.author or 'Unknown',
                'created': str(doc.core_properties.created or ''),
                'modified': str(doc.core_properties.modified or '')
            }
            
            current_section = {'title': None, 'content': []}
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if not text:
                    continue
                    
                # Check if paragraph is a heading
                if paragraph.style.name.startswith('Heading'):
                    if current_section['title'] and current_section['content']:
                        content['structure']['sections'].append(current_section.copy())
                    current_section = {'title': text, 'content': []}
                else:
                    current_section['content'].append(text)
                    content['text'].append(text)
                    
            # Add the last section if it exists
            if current_section['title'] and current_section['content']:
                content['structure']['sections'].append(current_section)
                
            # Extract tables if they exist
            content['tables'] = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                content['tables'].append(table_data)
                
            return content
            
        except Exception as e:
            raise DocumentProcessingError(f"Error processing DOCX: {str(e)}")

    def validate_content(self, content: Dict[str, Any]) -> bool:
        """Validate DOCX content"""
        return bool(content.get('text') or content.get('tables'))

class MarkdownProcessor(BaseDocumentProcessor):
    def extract_content(self, file_obj: BinaryIO) -> Dict[str, Any]:
        """Extract content from Markdown files"""
        try:
            text = file_obj.read().decode('utf-8')
            html = markdown.markdown(text, extensions=['tables', 'fenced_code'])
            soup = bs4.BeautifulSoup(html, 'html.parser')
            
            content = {
                'text': [],
                'metadata': {},
                'structure': {'sections': []},
                'type': 'markdown'
            }
            
            current_section = {'title': None, 'content': []}
            
            # Process headers and content
            for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'code', 'pre']):
                text = element.get_text().strip()
                if not text:
                    continue
                    
                if element.name.startswith('h'):
                    if current_section['title'] and current_section['content']:
                        content['structure']['sections'].append(current_section.copy())
                    current_section = {'title': text, 'content': []}
                else:
                    current_section['content'].append(text)
                    content['text'].append(text)
                    
            # Add the last section if it exists
            if current_section['title'] and current_section['content']:
                content['structure']['sections'].append(current_section)
                
            # Extract code blocks
            content['code_blocks'] = []
            for code_block in soup.find_all('pre'):
                content['code_blocks'].append({
                    'language': code_block.find('code').get('class', [''])[0] if code_block.find('code') else '',
                    'code': code_block.get_text().strip()
                })
                
            return content
            
        except Exception as e:
            raise DocumentProcessingError(f"Error processing Markdown: {str(e)}")

    def validate_content(self, content: Dict[str, Any]) -> bool:
        """Validate Markdown content"""
        return bool(content.get('text'))

class DocumentProcessor:
    def __init__(self):
        self.processors = {
            'pdf': PDFProcessor(),
            'docx': DocxProcessor(),
            'md': MarkdownProcessor()
        }
        
    def process_document(self, file_obj: BinaryIO, file_type: str) -> Dict[str, Any]:
        """Process document and return standardized content"""
        if file_type not in self.processors:
            raise UnsupportedFormatError(f"Format {file_type} not supported")
            
        processor = self.processors[file_type]
        content = processor.extract_content(file_obj)
        
        if not processor.validate_content(content):
            raise DocumentProcessingError("Invalid or empty document content")
            
        return processor.standardize_content(content)

    def detect_file_type(self, file_obj) -> str:
        """Detect file type from file object"""
        filename = getattr(file_obj, 'name', '').lower()
        if filename.endswith('.pdf'):
            return 'pdf'
        elif filename.endswith('.docx'):
            return 'docx'
        elif filename.endswith('.md'):
            return 'md'
        else:
            raise UnsupportedFormatError("Unsupported file type")

def process_uploaded_file(file_obj) -> Dict[str, Any]:
    """Process an uploaded file and return standardized content"""
    try:
        processor = DocumentProcessor()
        file_type = processor.detect_file_type(file_obj)
        return processor.process_document(file_obj, file_type)
    except Exception as e:
        st.error(f"Error processing file: {str(e)}")
        raise

# Content Analysis and Tutorial Generation
class ContentAnalyzer:
    def __init__(self, model):
        self.model = model
        
    def analyze_content(self, content: Dict[str, Any]) -> List[Topic]:
        """Analyze content and create topic structure"""
        try:
            # Extract the actual text content
            text_content = ' '.join(content['text']) if isinstance(content['text'], list) else content['text']
            
            # Get document structure
            sections = content.get('structure', {}).get('sections', [])
            
            if not text_content or not sections:
                raise ValueError("No content or sections found in document")

            # Create a more specific prompt for the AI
            prompt = f"""
            Analyze this educational content and create a structured tutorial outline.
            The content is about: {text_content[:200]}...

            Document sections:
            {json.dumps(sections, indent=2)}

            Create a detailed learning structure with the following JSON format:
            {{
                "lessons": [
                    {{
                        "title": "Specific topic title from the content",
                        "content": "Detailed explanation using actual content from the document",
                        "key_points": ["Key points extracted from the document"],
                        "practice": ["Practice items based on the content"],
                        "difficulty": "beginner|intermediate|advanced based on content complexity"
                    }}
                ]
            }}

            Base your response entirely on the provided document content.
            Each lesson should correspond to major sections or concepts from the document.
            """

            # Get AI response
            response = self.model.generate_content(prompt)
            
            try:
                structure = json.loads(response.text)
            except json.JSONDecodeError as e:
                # If JSON parsing fails, try to extract JSON from the response
                import re
                match = re.search(r'\{[\s\S]*\}', response.text)
                if match:
                    structure = json.loads(match.group())
                else:
                    raise ValueError(f"Could not parse AI response as JSON: {e}")

            # Convert AI response to topics
            topics = []
            for lesson in structure.get('lessons', []):
                # Create topic with actual content from the document
                topic = Topic(
                    title=lesson.get('title', 'Untitled Topic'),
                    content=lesson.get('content', ''),
                    subtopics=[],  # Initialize empty subtopics list
                    metadata={
                        'key_points': lesson.get('key_points', []),
                        'practice': lesson.get('practice', []),
                        'difficulty': lesson.get('difficulty', 'intermediate'),
                        'content_type': self._detect_content_type(content),
                        'source_sections': self._find_relevant_sections(lesson.get('title', ''), sections)
                    }
                )
                topics.append(topic)

            # If no topics were created, raise an exception
            if not topics:
                raise ValueError("No topics could be generated from the content")

            return topics

        except Exception as e:
            st.warning(f"Error analyzing content: {str(e)}")
            # Fall back to basic topic structure if AI analysis fails
            return self._create_basic_topics_from_sections(content)

    def _detect_content_type(self, content: Dict[str, Any]) -> str:
        """Detect the type of content based on its characteristics"""
        # Check for code-heavy content
        code_indicators = ['def ', 'class ', 'import ', 'return ', 'function', 'method']
        text_content = ' '.join(content['text']) if isinstance(content['text'], list) else content['text']
        
        # Count code-related keywords
        code_keyword_count = sum(1 for indicator in code_indicators if indicator in text_content.lower())
        
        # Check for code blocks in the content
        has_code_blocks = bool(content.get('code_blocks', []))
        
        # Calculate the ratio of lines that might be code
        total_lines = len(text_content.split('\n'))
        code_like_lines = len([line for line in text_content.split('\n') 
                             if any(ind in line for ind in code_indicators)])
        code_ratio = code_like_lines / total_lines if total_lines > 0 else 0
        
        # Decision logic for content type
        if has_code_blocks or code_ratio > 0.2 or code_keyword_count > 5:
            return 'technical'
        elif any(keyword in text_content.lower() 
                for keyword in ['theory', 'concept', 'principle', 'define', 'explain']):
            return 'theoretical'
        else:
            return 'practical'

    def _find_relevant_sections(self, topic_title: str, sections: List[Dict]) -> List[Dict]:
        """Find sections that are relevant to a topic title"""
        relevant_sections = []
        for section in sections:
            # Check if section title is similar to topic title
            if (section.get('title') and 
                self._calculate_similarity(section['title'].lower(), topic_title.lower()) > 0.5):
                relevant_sections.append(section)
        return relevant_sections

    def _calculate_similarity(self, str1: str, str2: str) -> float:
        """Calculate simple string similarity"""
        # Convert strings to sets of words
        set1 = set(str1.split())
        set2 = set(str2.split())
        
        # Calculate Jaccard similarity
        intersection = len(set1.intersection(set2))
        union = len(set1.union(set2))
        
        return intersection / union if union > 0 else 0

    def _create_basic_topics_from_sections(self, content: Dict[str, Any]) -> List[Topic]:
        """Create basic topics from document sections when AI analysis fails"""
        topics = []
        sections = content.get('structure', {}).get('sections', [])
        
        for section in sections:
            if section.get('title') and section.get('content'):
                topic = Topic(
                    title=section['title'],
                    content='\n'.join(section['content']) if isinstance(section['content'], list) 
                           else section['content'],
                    subtopics=[],
                    metadata={
                        'key_points': self._extract_key_points(section['content']),
                        'practice': [f"Explain the concept of {section['title']}"],
                        'difficulty': 'intermediate',
                        'content_type': self._detect_content_type(content)
                    }
                )
                topics.append(topic)
        
        return topics if topics else self._create_fallback_structure()

    def _extract_key_points(self, content: Union[str, List[str]]) -> List[str]:
        """Extract potential key points from content"""
        if isinstance(content, list):
            content = ' '.join(content)
            
        # Split into sentences
        sentences = content.split('.')
        
        # Look for sentences that might be key points
        key_points = []
        indicators = ['important', 'key', 'essential', 'fundamental', 'crucial', 'primary']
        
        for sentence in sentences:
            sentence = sentence.strip()
            if any(indicator in sentence.lower() for indicator in indicators) or \
               (len(sentence.split()) < 20 and sentence):  # Short, meaningful sentences
                key_points.append(sentence)
                
        return key_points[:5]  # Return up to 5 key points

    def _create_fallback_structure(self) -> List[Topic]:
        """Create a basic fallback topic structure"""
        return [Topic(
            title="Introduction",
            content="Basic overview of the content.",
            subtopics=[],
            metadata={
                'key_points': ['Basic understanding required'],
                'practice': ['Review the main concepts'],
                'difficulty': 'beginner',
                'content_type': 'practical'
            }
        )]

# Tutorial Generation Templates
class TutorialTemplate(ABC):
    @abstractmethod
    def create_tutorial_content(self, topic: Topic, user_performance: float) -> Dict[str, Any]:
        pass

class TechnicalTemplate(TutorialTemplate):
    def create_tutorial_content(self, topic: Topic, user_performance: float) -> Dict[str, Any]:
        """Template for technical content with dynamic handling based on document content"""
        difficulty = self._adjust_difficulty(user_performance)
        
        # Extract code examples and practice exercises
        practice_exercises = []
        if isinstance(topic.metadata.get('practice'), list):
            practice_exercises = [
                {
                    'title': f'Practice Exercise {i+1}',
                    'description': exercise
                }
                for i, exercise in enumerate(topic.metadata.get('practice', []))
            ]
        
        # Create the content structure
        content = f"""
# {topic.title}

## Overview 📋
{topic.content}

## Learning Objectives 🎯
{"".join(f"- {obj}\n" for obj in self._generate_objectives(topic, difficulty))}

## Detailed Explanation 📝
{self._generate_explanation(topic, difficulty)}
"""
        
        # Add practice exercises section if available
        if practice_exercises:
            content += "\n## Practice Exercises ✍️\n"
            for exercise in practice_exercises:
                content += f"""
### {exercise['title']}
{exercise['description']}
"""

        return {
            "content": content,
            "practice_exercises": practice_exercises,
            "metadata": {
                "difficulty": difficulty,
                "topic_type": "technical",
                "estimated_time": self._estimate_time(topic, difficulty)
            }
        }

    def _generate_objectives(self, topic: Topic, difficulty: str) -> List[str]:
        """Generate learning objectives based on topic content and difficulty"""
        objectives = []
        key_points = topic.metadata.get('key_points', [])
        
        if key_points:
            # Use existing key points as basis for objectives
            for point in key_points:
                if difficulty == "advanced":
                    objectives.append(f"Master {point}")
                elif difficulty == "intermediate":
                    objectives.append(f"Understand and apply {point}")
                else:
                    objectives.append(f"Learn basics of {point}")
        else:
            # Generate basic objectives from topic title and content
            objectives = [
                f"Understand core concepts of {topic.title}",
                "Apply learned concepts to practical problems",
                "Develop problem-solving skills in this area"
            ]
            
        return objectives

    def _generate_explanation(self, topic: Topic, difficulty: str) -> str:
        """Generate detailed explanation based on topic content and difficulty level"""
        base_explanation = topic.content
        
        # Add difficulty-specific additional content if available
        if difficulty == "advanced" and topic.metadata.get('advanced_content'):
            return f"""
{base_explanation}

### Advanced Concepts
{topic.metadata['advanced_content']}
"""
        elif difficulty == "intermediate" and topic.metadata.get('intermediate_content'):
            return f"""
{base_explanation}

### Key Concepts
{topic.metadata['intermediate_content']}
"""
        else:
            return base_explanation

    def _estimate_time(self, topic: Topic, difficulty: str) -> str:
        """Estimate time needed based on content complexity and difficulty"""
        base_time = 30  # base time in minutes
        
        # Adjust for content length
        content_length = len(topic.content.split())
        time_factor = content_length / 500  # adjust time based on content length
        
        # Adjust for difficulty
        difficulty_multiplier = {
            "beginner": 1,
            "intermediate": 1.5,
            "advanced": 2
        }
        
        estimated_time = base_time * time_factor * difficulty_multiplier.get(difficulty, 1)
        return f"{int(estimated_time)} minutes"

    def _adjust_difficulty(self, user_performance: float) -> str:
        """Adjust difficulty based on user performance"""
        if user_performance >= 85:
            return "advanced"
        elif user_performance >= 65:
            return "intermediate"
        return "beginner"

class TheoreticalTemplate(TutorialTemplate):
    def create_tutorial_content(self, topic: Topic, user_performance: float) -> Dict[str, Any]:
        """Create tutorial content based on theoretical approach"""
        # Previous incorrect create_tutorial method has been removed
        difficulty = self._adjust_difficulty(user_performance)
        return {
            "content": {
                "title": topic.title,
                "theory": self._generate_theory(topic),
                "concepts": self._generate_concepts(topic, difficulty),
                "examples": self._generate_examples(topic, difficulty)
            },
            "difficulty": difficulty
        }
    
    def _adjust_difficulty(self, user_performance: float) -> str:
        if user_performance >= 85:
            return "advanced"
        elif user_performance >= 65:
            return "intermediate"
        return "beginner"

class PracticalTemplate(TutorialTemplate):
    def create_tutorial_content(self, topic: Topic, user_performance: float) -> Dict[str, Any]:
        difficulty = self._adjust_difficulty(user_performance)
        return {
            "overview": {
                "title": topic.title,
                "practical_context": self._generate_context(topic),
                "learning_goals": self._generate_goals(topic, difficulty)
            },
            "content": {
                "steps": self._generate_steps(topic),
                "examples": self._generate_examples(topic, difficulty),
                "exercises": self._generate_exercises(topic, difficulty)
            },
            "difficulty": difficulty
        }

class AdaptiveTutorialGenerator:
    def __init__(self):
        self.templates = {
            'technical': TechnicalTemplate(),
            'theoretical': TheoreticalTemplate(),
            'practical': PracticalTemplate()
        }
        
    def generate_tutorial(self, topic: Topic, user_performance: float) -> Dict[str, Any]:
        """Generate tutorial content based on topic type and user performance"""
        template = self.templates.get(topic.content_type, TechnicalTemplate())
        return template.create_tutorial_content(topic, user_performance)

    def adapt_difficulty(self, user_performance: float, current_difficulty: str) -> str:
        """Adapt difficulty based on user performance"""
        if user_performance >= 85 and current_difficulty != "advanced":
            return "advanced"
        elif user_performance >= 65 and current_difficulty == "beginner":
            return "intermediate"
        elif user_performance < 65 and current_difficulty != "beginner":
            return "beginner"
        return current_difficulty

    def generate_alternative_explanation(self, topic: Topic, 
                                      previous_approach: str) -> Dict[str, Any]:
        """Generate alternative explanation when user struggles"""
        # Define possible approaches
        approaches = ["practical", "theoretical", "technical"]
        
        # Choose next approach (different from previous)
        next_approach = next(
            (a for a in approaches if a != previous_approach), 
            "practical"
        )
        
        # Use corresponding template for new approach
        template = self.templates[next_approach]
        
        # Generate content with modified difficulty
        alternative_content = template.create_tutorial_content(topic, 50.0)  # Start at middle difficulty
        
        # Add meta information about the alternative approach
        alternative_content["meta"] = {
            "approach": next_approach,
            "reason": "Alternative explanation provided due to learning difficulty",
            "focus": self._get_approach_focus(next_approach)
        }
        
        return alternative_content

    def _get_approach_focus(self, approach: str) -> str:
        """Get the focus area for each approach type"""
        focus_map = {
            "practical": "hands-on examples and real-world applications",
            "theoretical": "underlying concepts and principles",
            "technical": "implementation details and specific techniques"
        }
        return focus_map.get(approach, "general understanding")

class TutorialManager:
    """Enhanced Tutorial Manager with improved evaluation handling"""
    
    def __init__(self, model):
        self.content_analyzer = ContentAnalyzer(model)
        self.tutorial_generator = AdaptiveTutorialGenerator()
        self.evaluation_engine = EnhancedEvaluationEngine(model)  # Use new evaluation engine
        self.current_performance = {
            'understanding_level': 50.0,
            'cognitive_skills': {
                'comprehension': 50.0,
                'application': 50.0,
                'analysis': 50.0
            },
            'response_quality': {
                'clarity': 50.0,
                'completeness': 50.0,
                'accuracy': 50.0,
                'depth': 50.0
            }
        }
        self.approach_history = []
        self.topic_evaluations = {}  # Store detailed evaluations per topic
        
    def create_tutorial(self, content: Dict[str, Any]) -> List[Topic]:
        """Create a new tutorial from content"""
        topics = self.content_analyzer.analyze_content(content)
        # Initialize evaluation storage for each topic
        for topic in topics:
            self.topic_evaluations[topic.title] = []
        return topics
        
    def generate_next_content(self, topic: Topic) -> Dict[str, Any]:
        """Generate next piece of tutorial content with adaptive difficulty"""
        # Use cognitive skills to adjust content difficulty
        cognitive_level = sum(self.current_performance['cognitive_skills'].values()) / 3
        
        # Adjust content based on response quality
        quality_level = sum(self.current_performance['response_quality'].values()) / 4
        
        # Generate tutorial content with enhanced parameters
        return self.tutorial_generator.generate_tutorial(
            topic=topic,
            user_performance=self.current_performance['understanding_level'],
            cognitive_level=cognitive_level,
            quality_level=quality_level,
            previous_evaluations=self.topic_evaluations.get(topic.title, [])
        )
        
    def evaluate_response(self, user_response: str, topic: Topic) -> Dict[str, Any]:
        """Evaluate user response with enhanced feedback"""
        try:
            # Get expected points from topic metadata
            expected_points = topic.metadata.get('key_points', []) if topic.metadata else []
            
            # Use enhanced evaluation engine
            evaluation_result = self.evaluation_engine.evaluate_response(
                user_response=user_response,
                expected_points=expected_points,
                topic=topic
            )
            
            # Store evaluation for topic history
            self.topic_evaluations[topic.title].append({
                'response': user_response,
                'evaluation': evaluation_result,
                'timestamp': time.time()
            })
            
            # Update performance metrics
            self._update_performance_metrics(evaluation_result)
            
            # Generate adaptive recommendations
            recommendations = self._generate_adaptive_recommendations(
                topic=topic,
                evaluation_result=evaluation_result
            )
            
            # Add recommendations to evaluation result
            evaluation_result['recommendations'] = recommendations
            
            return evaluation_result
            
        except Exception as e:
            st.error(f"Error in evaluation: {str(e)}")
            return self._create_detailed_fallback_evaluation(topic)
            
    def _update_performance_metrics(self, evaluation_result: Dict[str, Any]) -> None:
        """Update all performance metrics from evaluation"""
        # Update understanding level
        self.current_performance['understanding_level'] = evaluation_result.get(
            'understanding_level',
            self.current_performance['understanding_level']
        )
        
        # Update cognitive skills
        eval_cognitive = evaluation_result.get('evaluation_data', {}).get('cognitive_skills', {})
        for skill, value in eval_cognitive.items():
            if skill in self.current_performance['cognitive_skills']:
                # Use exponential moving average for smooth transitions
                alpha = 0.3  # Smoothing factor
                current = self.current_performance['cognitive_skills'][skill]
                self.current_performance['cognitive_skills'][skill] = (
                    alpha * value + (1 - alpha) * current
                )
        
        # Update response quality metrics
        eval_quality = evaluation_result.get('evaluation_data', {}).get('response_quality', {})
        for metric, value in eval_quality.items():
            if metric in self.current_performance['response_quality']:
                # Use exponential moving average
                alpha = 0.3
                current = self.current_performance['response_quality'][metric]
                self.current_performance['response_quality'][metric] = (
                    alpha * value + (1 - alpha) * current
                )
    
    def _generate_adaptive_recommendations(self, 
                                        topic: Topic,
                                        evaluation_result: Dict[str, Any]) -> List[str]:
        """Generate personalized recommendations based on performance history"""
        recommendations = []
        
        # Analyze performance trends
        topic_history = self.topic_evaluations[topic.title]
        if len(topic_history) >= 2:
            # Check for recurring issues
            recurring_issues = self._identify_recurring_issues(topic_history)
            for issue in recurring_issues:
                recommendations.append(f"Focus on improving {issue}")
        
        # Add cognitive skill-based recommendations
        cognitive_skills = self.current_performance['cognitive_skills']
        lowest_skill = min(cognitive_skills.items(), key=lambda x: x[1])
        if lowest_skill[1] < 70:
            recommendations.append(
                f"Work on improving {lowest_skill[0]} through additional practice"
            )
        
        # Add quality-based recommendations
        quality_metrics = self.current_performance['response_quality']
        lowest_quality = min(quality_metrics.items(), key=lambda x: x[1])
        if lowest_quality[1] < 70:
            recommendations.append(
                f"Focus on improving {lowest_quality[0]} in your responses"
            )
        
        # Add topic-specific recommendations
        topic_recommendations = self._get_topic_specific_recommendations(
            topic,
            evaluation_result
        )
        recommendations.extend(topic_recommendations)
        
        return recommendations
    
    def _identify_recurring_issues(self, topic_history: List[Dict]) -> List[str]:
        """Identify recurring issues in user responses"""
        issues = []
        
        # Analyze last 3 evaluations
        recent_evals = topic_history[-3:]
        
        # Check for consistently low scores in specific areas
        areas = {
            'comprehension': [],
            'application': [],
            'analysis': [],
            'clarity': [],
            'completeness': [],
            'accuracy': [],
            'depth': []
        }
        
        for eval_data in recent_evals:
            evaluation = eval_data['evaluation']
            
            # Collect cognitive skills scores
            cognitive = evaluation.get('evaluation_data', {}).get('cognitive_skills', {})
            for skill, score in cognitive.items():
                areas[skill].append(score)
            
            # Collect quality metrics
            quality = evaluation.get('evaluation_data', {}).get('response_quality', {})
            for metric, score in quality.items():
                areas[metric].append(score)
        
        # Identify consistently low areas
        for area, scores in areas.items():
            if scores and all(score < 70 for score in scores):
                issues.append(area)
        
        return issues
    
    def _get_topic_specific_recommendations(self, 
                                          topic: Topic,
                                          evaluation_result: Dict[str, Any]) -> List[str]:
        """Generate recommendations specific to the topic"""
        recommendations = []
        
        # Get missing points
        point_analysis = evaluation_result.get('evaluation_data', {}).get('point_analysis', [])
        missing_points = [
            point['point'] for point in point_analysis 
            if point['coverage'] == 'none'
        ]
        
        if missing_points:
            recommendations.append(
                f"Review these key points: {', '.join(missing_points)}"
            )
        
        # Add specific practice suggestions based on topic type
        if topic.metadata and topic.metadata.get('content_type') == 'technical':
            recommendations.append("Practice with hands-on examples")
        elif topic.metadata and topic.metadata.get('content_type') == 'theoretical':
            recommendations.append("Focus on understanding underlying principles")
        
        return recommendations
    
    def _create_detailed_fallback_evaluation(self, topic: Topic) -> Dict[str, Any]:
        """Create detailed fallback evaluation when main evaluation fails"""
        return {
            "understanding_level": 50.0,
            "feedback": f"""
## Feedback Analysis 📊

We couldn't generate detailed feedback for your response about {topic.title}.
Please try again and:
1. Address all key points
2. Provide specific examples
3. Explain concepts clearly

### Understanding Level
▰▰▰▰▰▱▱▱▱▱ 50%
""",
            "evaluation_data": {
                "cognitive_skills": {
                    "comprehension": 50.0,
                    "application": 50.0,
                    "analysis": 50.0
                },
                "response_quality": {
                    "clarity": 50.0,
                    "completeness": 50.0,
                    "accuracy": 50.0,
                    "depth": 50.0
                }
            },
            "recommendations": [
                "Please provide a more detailed response",
                "Include specific examples",
                "Address all key points from the topic"
            ]
        }

class EnhancedEvaluationEngine:
    """Enhanced evaluation engine with improved response analysis"""
    
    def __init__(self, model):
        self.model = model
        
    def evaluate_response(self, user_response: str, 
                         expected_points: List[str], 
                         topic: Topic) -> Dict[str, Any]:
        """Evaluate user response with enhanced analysis"""
        try:
            # Check for low-effort responses
            if self._is_low_effort_response(user_response):
                return self._generate_low_effort_feedback()
            
            # Generate evaluation prompt
            prompt = self._create_enhanced_evaluation_prompt(
                user_response,
                expected_points,
                topic
            )
            
            # Get AI evaluation
            response = self.model.generate_content(prompt)
            
            try:
                evaluation = json.loads(clean_json_string(response.text))
            except json.JSONDecodeError:
                # Fallback to structured extraction if JSON parsing fails
                evaluation = self._extract_structured_evaluation(response.text)
            
            # Enhance evaluation with detailed analysis
            enhanced_evaluation = self._enhance_evaluation(evaluation, user_response, expected_points)
            
            # Calculate nuanced understanding level
            understanding_level = self._calculate_nuanced_understanding(enhanced_evaluation)
            
            # Generate specific, actionable feedback
            feedback = self._generate_detailed_feedback(enhanced_evaluation)
            
            return {
                "understanding_level": understanding_level,
                "feedback": feedback,
                "evaluation_data": enhanced_evaluation,
                "recommendations": self._generate_targeted_recommendations(enhanced_evaluation)
            }
            
        except Exception as e:
            st.error(f"Error in evaluation: {str(e)}")
            return self._create_detailed_fallback_evaluation()
    
    def _is_low_effort_response(self, response: str) -> bool:
        """Detect low-effort or non-serious responses"""
        response = response.lower().strip()
        
        # Check response length
        if len(response.split()) < 3:
            return True
            
        # Check for common low-effort responses
        low_effort_patterns = [
            r'^(ok|cool|nice|good|yes|no|maybe)[\s\.]*$',
            r'^(i understand|got it|makes sense)[\s\.]*$',
            r'^(idk|dunno|whatever)[\s\.]*$'
        ]
        
        return any(re.match(pattern, response) for pattern in low_effort_patterns)
    
    def _create_enhanced_evaluation_prompt(self, user_response: str,
                                         expected_points: List[str],
                                         topic: Topic) -> str:
        """Create a more nuanced evaluation prompt"""
        return f"""
        Evaluate this response about {topic.title} using the following criteria:

        User response: "{user_response}"
        
        Expected key points:
        {json.dumps(expected_points, indent=2)}
        
        Evaluate and return a JSON object with:
        {{
            "point_analysis": [
                {{
                    "point": "expected point",
                    "coverage": "full|partial|none",
                    "understanding": "deep|basic|superficial|incorrect",
                    "evidence": "relevant text from response",
                    "improvement_needed": "specific improvement suggestion"
                }}
            ],
            "cognitive_skills": {{
                "comprehension": 0-100,
                "application": 0-100,
                "analysis": 0-100
            }},
            "response_quality": {{
                "clarity": 0-100,
                "completeness": 0-100,
                "accuracy": 0-100,
                "depth": 0-100
            }},
            "misconceptions": [
                {{
                    "identified": "misconception",
                    "correction": "correct understanding",
                    "explanation": "why it's wrong"
                }}
            ],
            "strengths": [
                {{
                    "aspect": "what was done well",
                    "example": "evidence from response"
                }}
            ],
            "improvement_areas": [
                {{
                    "aspect": "what needs improvement",
                    "current_state": "current understanding",
                    "target_state": "desired understanding",
                    "suggestion": "specific improvement action"
                }}
            ]
        }}

        Base evaluation on:
        1. Accuracy and completeness of point coverage
        2. Depth of understanding demonstrated
        3. Application of concepts
        4. Clarity of explanation
        5. Identification of misconceptions
        
        Provide specific evidence and examples from the response.
        """
    
    def _calculate_nuanced_understanding(self, evaluation: Dict[str, Any]) -> float:
        """Calculate understanding level with weighted criteria"""
        weights = {
            'point_coverage': 0.35,
            'cognitive_skills': 0.25,
            'response_quality': 0.25,
            'misconceptions': 0.15
        }
        
        # Calculate point coverage score
        point_scores = {
            'full': 1.0,
            'partial': 0.5,
            'none': 0.0
        }
        
        point_coverage = sum(point_scores[point['coverage']] 
                           for point in evaluation.get('point_analysis', [])) / \
                        len(evaluation.get('point_analysis', [1]))
        
        # Calculate cognitive skills score
        cognitive_skills = sum(evaluation.get('cognitive_skills', {}).values()) / 300
        
        # Calculate response quality score
        response_quality = sum(evaluation.get('response_quality', {}).values()) / 400
        
        # Calculate misconceptions impact
        misconceptions_count = len(evaluation.get('misconceptions', []))
        misconceptions_impact = max(0, 1 - (misconceptions_count * 0.2))
        
        # Calculate weighted final score
        understanding_level = (
            point_coverage * weights['point_coverage'] +
            cognitive_skills * weights['cognitive_skills'] +
            response_quality * weights['response_quality'] +
            misconceptions_impact * weights['misconceptions']
        ) * 100
        
        return round(max(0, min(100, understanding_level)), 1)
    
    def _generate_detailed_feedback(self, evaluation: Dict[str, Any]) -> str:
        """Generate specific, actionable feedback"""
        # Format point-by-point analysis
        point_analysis = "\n".join([
            f"### {point['point']}\n"
            f"- Coverage: {point['coverage'].title()}\n"
            f"- Understanding: {point['understanding'].title()}\n"
            f"- Evidence: \"{point['evidence']}\"\n"
            f"- To improve: {point['improvement_needed']}\n"
            for point in evaluation.get('point_analysis', [])
        ])
        
        # Format cognitive skills feedback
        cognitive_skills = evaluation.get('cognitive_skills', {})
        skills_feedback = (
            f"- Comprehension: {'▰' * int(cognitive_skills.get('comprehension', 0)/10)}"
            f"{'▱' * (10-int(cognitive_skills.get('comprehension', 0)/10))}\n"
            f"- Application: {'▰' * int(cognitive_skills.get('application', 0)/10)}"
            f"{'▱' * (10-int(cognitive_skills.get('application', 0)/10))}\n"
            f"- Analysis: {'▰' * int(cognitive_skills.get('analysis', 0)/10)}"
            f"{'▱' * (10-int(cognitive_skills.get('analysis', 0)/10))}\n"
        )
        
        return f"""
## Detailed Feedback Analysis 📊

### Point-by-Point Analysis
{point_analysis}

### Cognitive Skills Assessment
{skills_feedback}

### Key Strengths 💪
{"".join(f"- {strength['aspect']}\n  Example: \"{strength['example']}\"\n" for strength in evaluation.get('strengths', []))}

### Areas for Improvement 📈
{"".join(f"- {area['aspect']}\n  Current: {area['current_state']}\n  Goal: {area['target_state']}\n  Suggestion: {area['suggestion']}\n" for area in evaluation.get('improvement_areas', []))}

### Understanding Level
{'▰' * int(evaluation.get('understanding_level', 0)/10)}{'▱' * (10-int(evaluation.get('understanding_level', 0)/10))} {evaluation.get('understanding_level', 0)}%
"""
    
    def _generate_targeted_recommendations(self, evaluation: Dict[str, Any]) -> List[str]:
        """Generate specific, actionable recommendations"""
        recommendations = []
        
        # Add recommendations based on cognitive skills
        cognitive_skills = evaluation.get('cognitive_skills', {})
        if cognitive_skills.get('comprehension', 0) < 70:
            recommendations.append("Review the core concepts and definitions")
        if cognitive_skills.get('application', 0) < 70:
            recommendations.append("Practice applying concepts to real-world scenarios")
        if cognitive_skills.get('analysis', 0) < 70:
            recommendations.append("Work on analyzing relationships between concepts")
        
        # Add recommendations based on response quality
        quality = evaluation.get('response_quality', {})
        if quality.get('clarity', 0) < 70:
            recommendations.append("Focus on explaining concepts more clearly and systematically")
        if quality.get('completeness', 0) < 70:
            recommendations.append("Ensure all key points are addressed in your responses")
        if quality.get('depth', 0) < 70:
            recommendations.append("Provide more detailed explanations with specific examples")
        
        # Add specific improvement suggestions
        for area in evaluation.get('improvement_areas', []):
            recommendations.append(area['suggestion'])
        
        return recommendations
    
    def _generate_low_effort_feedback(self) -> Dict[str, Any]:
        """Generate feedback for low-effort responses"""
        return {
            "understanding_level": 0,
            "feedback": """
## Feedback Analysis 📊

Your response appears to be too brief or general to demonstrate understanding of the topic.
Please provide a more detailed response that:
1. Addresses the key points
2. Demonstrates your understanding
3. Includes specific examples or explanations

### Understanding Level
▱▱▱▱▱▱▱▱▱▱ 0%
""",
            "evaluation_data": {
                "point_analysis": [],
                "cognitive_skills": {
                    "comprehension": 0,
                    "application": 0,
                    "analysis": 0
                },
                "response_quality": {
                    "clarity": 0,
                    "completeness": 0,
                    "accuracy": 0,
                    "depth": 0
                }
            },
            "recommendations": [
                "Provide a complete response that addresses all key points",
                "Demonstrate your understanding with specific examples",
                "Explain concepts in your own words"
            ]
        }
# Main Application Logic and Streamlit Interface

def init_gemini(api_key: str = None) -> Optional[Any]:
    """Initialize the Gemini AI model"""
    try:
        if api_key:
            genai.configure(api_key=api_key)
            return genai.GenerativeModel('gemini-pro')
        return None
    except Exception as e:
        st.error(f"Error initializing Gemini model: {str(e)}")
        return None

def init_session_state():
    """Initialize or reset session state"""
    if 'tutorial_manager' not in st.session_state:
        st.session_state.tutorial_manager = None
    if 'current_topic_index' not in st.session_state:
        st.session_state.current_topic_index = 0
    if 'topics' not in st.session_state:
        st.session_state.topics = []
    if 'conversation_history' not in st.session_state:
        st.session_state.conversation_history = []
    if 'user_progress' not in st.session_state:
        st.session_state.user_progress = {
            'understanding_level': 50.0,
            'completed_topics': [],
            'current_approach': 'practical'
        }

def render_sidebar():
    """Render the application sidebar"""
    with st.sidebar:
        st.markdown("## 📚 Learning Progress")
        
        if st.session_state.topics:
            # Display progress bar
            completed = len(st.session_state.user_progress['completed_topics'])
            total = len(st.session_state.topics)
            progress = int((completed / total) * 100)
            st.progress(progress)
            st.markdown(f"**{progress}%** completed ({completed}/{total} topics)")
            
            # Display topic list
            st.markdown("### Topics Overview")
            for i, topic in enumerate(st.session_state.topics):
                status = "✅" if topic.title in st.session_state.user_progress['completed_topics'] else \
                         "📍" if i == st.session_state.current_topic_index else "⭕️"
                st.markdown(
                    f"<div class='topic-item {status}'>{i+1}. {topic.title}</div>",
                    unsafe_allow_html=True
                )
            
            # Display current stats
            st.markdown("### Current Stats")
            st.markdown(f"Understanding Level: {st.session_state.user_progress['understanding_level']:.1f}%")
            st.markdown(f"Learning Approach: {st.session_state.user_progress['current_approach'].title()}")
            
            # Reset button
            if st.button("🔄 Reset Tutorial"):
                reset_session()
                st.rerun()

def render_header():
    """Render the application header"""
    st.markdown("""
        <div style='text-align: center; padding: 2rem 0;'>
            <h1>🎓 AI Learning Assistant</h1>
            <p style='font-size: 1.2rem; color: #4B5563;'>
                Your personalized learning journey with AI guidance
            </p>
        </div>
    """, unsafe_allow_html=True)

def render_file_upload():
    """Render the file upload section"""
    st.markdown("""
        <div class='content-block'>
            <h2>📚 Upload Learning Material</h2>
            <p>Upload your educational content to begin the learning journey.</p>
        </div>
    """, unsafe_allow_html=True)
    
    uploaded_file = st.file_uploader(
        "Upload Educational Content",
        type=["pdf", "docx", "md"],
        help="Supported formats: PDF, Word documents, and Markdown files"
    )
    
    return uploaded_file

def render_chat_interface():
    """Render the chat interface"""
    if not st.session_state.topics:
        return
        
    # Get current topic
    current_topic = st.session_state.topics[st.session_state.current_topic_index]
    
    # Display conversation history
    for message in st.session_state.conversation_history:
        with st.chat_message(message["role"]):
            st.markdown(message["content"], unsafe_allow_html=True)
    
    # Generate new content if needed
    if not st.session_state.conversation_history or \
       st.session_state.conversation_history[-1]["role"] == "user":
        tutorial_content = st.session_state.tutorial_manager.generate_next_content(current_topic)
        
        with st.chat_message("assistant"):
            st.markdown(tutorial_content["content"], unsafe_allow_html=True)
        
        st.session_state.conversation_history.append({
            "role": "assistant",
            "content": tutorial_content["content"]
        })
    
    # Handle user input
    user_input = st.chat_input("Your response...")
    if user_input:
        # Display user message
        with st.chat_message("user"):
            st.markdown(user_input, unsafe_allow_html=True)
        
        st.session_state.conversation_history.append({
            "role": "user",
            "content": user_input
        })
        
        # Evaluate response
        evaluation = st.session_state.tutorial_manager.evaluate_response(
            user_input,
            current_topic
        )
        
        # Display feedback
        with st.chat_message("assistant"):
            st.markdown(evaluation["feedback"], unsafe_allow_html=True)
        
        st.session_state.conversation_history.append({
            "role": "assistant",
            "content": evaluation["feedback"]
        })
        
        # Update progress
        st.session_state.user_progress['understanding_level'] = evaluation["understanding_level"]
        
        # Check if topic is completed
        if evaluation["understanding_level"] >= 70:
            st.session_state.user_progress['completed_topics'].append(current_topic.title)
            advance_topic()
            st.rerun()
            
def handle_file_upload(uploaded_file):
    try:
        with st.spinner("🔄 Processing content..."):
            content = process_uploaded_file(uploaded_file)
            
            if not content or not content.get('text'):
                st.error("No content could be extracted from the file")
                return
                
            if not st.session_state.tutorial_manager:
                st.session_state.tutorial_manager = TutorialManager(st.session_state.model)
            
            topics = st.session_state.tutorial_manager.create_tutorial(content)
            
            if not topics:
                st.error("Could not create tutorial from content")
                return
                
            st.session_state.topics = topics
            st.success("✅ Tutorial created successfully!")
            st.rerun()
            
    except Exception as e:
        st.error(f"❌ Error processing file: {str(e)}")

def advance_topic():
    """Advance to the next topic"""
    if st.session_state.current_topic_index < len(st.session_state.topics) - 1:
        st.session_state.current_topic_index += 1
        st.session_state.conversation_history = []
    else:
        show_completion_message()

def show_completion_message():
    """Show tutorial completion message"""
    st.balloons()
    st.success("""
        🎉 Congratulations! You've completed the tutorial!
        
        Here's your learning summary:
        - Topics completed: {} out of {}
        - Final understanding level: {:.1f}%
        - Learning approach: {}
        
        Would you like to:
        1. Start a new tutorial with different content
        2. Review specific topics
        3. Get a detailed progress report
    """.format(
        len(st.session_state.user_progress['completed_topics']),
        len(st.session_state.topics),
        st.session_state.user_progress['understanding_level'],
        st.session_state.user_progress['current_approach']
    ))

def reset_session():
    """Reset the session state"""
    st.session_state.tutorial_manager = None
    st.session_state.current_topic_index = 0
    st.session_state.topics = []
    st.session_state.conversation_history = []
    st.session_state.user_progress = {
        'understanding_level': 50.0,
        'completed_topics': [],
        'current_approach': 'practical'
    }

def apply_custom_css():
    """Apply custom CSS styling"""
    st.markdown("""
        <style>
        /* Global Styles */
        .main {
            padding: 1rem;
            max-width: 1200px;
            margin: 0 auto;
        }
        
        /* Typography */
        h1 {
            color: #1E3A8A;
            font-size: 2.5rem;
            margin-bottom: 1rem;
        }
        h2 {
            color: #2563EB;
            font-size: 1.8rem;
            margin-bottom: 0.8rem;
        }
        h3 {
            color: #3B82F6;
            font-size: 1.5rem;
            margin-bottom: 0.6rem;
        }
        
        /* Content Blocks */
        .content-block {
            background: white;
            padding: 1.5rem;
            border-radius: 8px;
            border: 1px solid #E5E7EB;
            margin-bottom: 1rem;
            box-shadow: 0 1px 3px rgba(0,0,0,0.1);
        }
        
        /* Chat Interface */
        .chat-message {
            padding: 1rem;
            border-radius: 8px;
            margin: 0.5rem 0;
            border: 1px solid #E5E7EB;
        }
        .chat-message.assistant {
            background: #F8FAFC;
        }
        .chat-message.user {
            background: #F0F9FF;
        }
        
        /* Topic List */
        .topic-item {
            padding: 0.75rem;
            margin: 0.25rem 0;
            border-radius: 4px;
            transition: background-color 0.2s;
        }
        .topic-item:hover {
            background: #F3F4F6;
        }
        .topic-item.✅ {
            color: #059669;
        }
        .topic-item.📍 {
            color: #2563EB;
            font-weight: bold;
        }
        
        /* Progress Bar */
        .stProgress .st-bo {
            height: 8px;
            border-radius: 4px;
        }
        
        /* Buttons */
        .stButton button {
            background: #2563EB;
            color: white;
            border: none;
            padding: 0.5rem 1rem;
            border-radius: 4px;
            font-weight: 500;
            transition: background-color 0.2s;
        }
        .stButton button:hover {
            background: #1D4ED8;
        }
        
        /* File Uploader */
        .stFileUploader {
            padding: 2rem;
            border: 2px dashed #E5E7EB;
            border-radius: 8px;
            text-align: center;
        }
        
        /* Feedback Sections */
        .feedback-section {
            margin: 1rem 0;
            padding: 1rem;
            border-left: 4px solid #2563EB;
            background: #F8FAFC;
        }
        </style>
    """, unsafe_allow_html=True)
def generate_teaching_message(topic: Topic, phase: str, conversation_history: List[Dict], model) -> dict:
    """Generate teaching message using adaptive tutorial generator"""
    try:
        # Get user performance from conversation history
        user_performance = calculate_user_performance(conversation_history)
        
        # Generate adaptive tutorial content
        generator = AdaptiveTutorialGenerator()
        tutorial_content = generator.generate_tutorial(topic, user_performance)
        
        # Format the message content
        formatted_content = f"""
# {topic.title} 📚

## Overview
{tutorial_content.get('content', 'No content available.')}

## Code Examples 💻
"""
        
        # Add code examples if available
        for example in tutorial_content.get('code_examples', []):
            formatted_content += f"""
### {example['title']}
```python
{example['code']}
```
"""

        # Add practice exercises
        formatted_content += "\n## Practice Exercises ✍️\n"
        for exercise in tutorial_content.get('practice_exercises', []):
            formatted_content += f"""
### {exercise['title']}
{exercise['description']}
"""

        return {
            "content": formatted_content,
            "examples": tutorial_content.get('code_examples', []),
            "exercises": tutorial_content.get('practice_exercises', []),
            "metadata": tutorial_content.get('metadata', {}),
            "key_points": topic.metadata.get('key_points', []) if topic.metadata else []
        }
        
    except Exception as e:
        st.error(f"Error generating teaching message: {str(e)}")
        return create_fallback_content(topic)

def calculate_user_performance(conversation_history: List[Dict]) -> float:
    """Calculate user performance from conversation history"""
    if not conversation_history:
        return 50.0
        
    # Look for evaluation results in recent messages
    for message in reversed(conversation_history):
        if message["role"] == "assistant" and "Understanding Level" in message["content"]:
            try:
                # Extract understanding level from the message
                level_text = message["content"].split("Understanding Level")[1]
                level = float(level_text.split("%")[0].strip())
                return level
            except:
                continue
                
    return 50.0  # Default value if no evaluation found

def create_fallback_content(topic: Topic) -> dict:
    """Create basic content structure when main generation fails"""
    fallback_content = f"""
# {topic.title} 📚

## Overview
{topic.content}

## Key Points
- Basic understanding of the topic is essential
- Practice and repetition help master the concepts
- Seek additional resources for deeper understanding

## Practice
Can you explain the main concepts of {topic.title}?
"""
    
    return {
        "content": fallback_content,
        "examples": [],
        "exercises": [{
            "title": "Basic Understanding",
            "description": f"Explain the key concepts of {topic.title}"
        }],
        "metadata": {"difficulty": "beginner"},
        "key_points": ["Basic understanding of concepts"]
    }
def main():
    """Main application function"""
    # Page configuration
    st.set_page_config(
        page_title="AI Learning Assistant 📚",
        page_icon="🎓",
        layout="wide",
        initial_sidebar_state="expanded"
    )
    
    # Apply custom CSS
    apply_custom_css()
    
    # Initialize session state
    init_session_state()
    
    # Get API key
    api_key = None
    try:
        api_key = st.secrets["GEMINI_API_KEY"]
    except:
        with st.expander("🔑 API Key Configuration"):
            api_key = st.text_input("Enter your Gemini API Key:", type="password")
    
    if not api_key:
        st.error("⚠️ Please provide your API key to continue")
        st.stop()
    
    # Initialize model
    if 'model' not in st.session_state:
        st.session_state.model = init_gemini(api_key)
    
    if not st.session_state.model:
        st.error("❌ Failed to initialize AI model")
        st.stop()
    
    # Render header
    render_header()
    
    # Create two-column layout
    col1, col2 = st.columns([7, 3])
    
    with col1:
        if not st.session_state.topics:
            uploaded_file = render_file_upload()
            if uploaded_file:
                handle_file_upload(uploaded_file)
        else:
            render_chat_interface()
    
    with col2:
        render_sidebar()

if __name__ == "__main__":
    main()
